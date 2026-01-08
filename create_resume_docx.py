#!/usr/bin/env python3
"""
Create resume.docx from resume content
Install python-docx: pip3 install python-docx
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Header Section
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run('ABUBAKAR HUSSAIN')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(5, 13, 24)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run('Senior Flutter & .NET Full-Stack Developer | AI-Powered Development Specialist')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    contact = doc.add_paragraph()
    contact.add_run('📧 abubakar-hussain@hotmail.com').font.size = Pt(10)
    contact.add_run('  |  📱 +92-300-1115589').font.size = Pt(10)
    contact.add_run('  |  🔗 LinkedIn: linkedin.com/in/abubakar-hussain-66b834130').font.size = Pt(10)
    contact.add_run('  |  🌐 Portfolio: info-abubakar-hussain.github.io/abubakarhussain').font.size = Pt(10)
    contact.add_run('  |  📍 Paragon City, Lahore, Punjab, Pakistan').font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    summary = doc.add_paragraph(
        'Results-driven Senior Mobile & Full-Stack Developer with 10+ years of software development experience. '
        'Expert in Flutter (5+ years), Android Native (4+ years), and .NET Backend Development. Specialized in building '
        'enterprise-grade applications with AI/ML integration, facial recognition (InsightFace), OCR document verification '
        '(Qwen2-7B), BioPass fingerprint biometric authentication, blockchain technology, and real-time communication systems '
        '(XMPP/Ejabberd, WebSocket/SignalR). AI-augmented development leader using Cursor AI, GitHub Copilot, and vibe coding '
        'techniques to accelerate delivery by 40% while maintaining 95% code quality. Proven track record in leading development '
        'teams, delivering high-quality cross-platform applications, and implementing cutting-edge solutions for mobile, web, and backend systems.'
    )
    summary_format = summary.paragraph_format
    summary_format.space_after = Pt(12)
    
    # Core Competencies
    doc.add_heading('CORE COMPETENCIES', level=1)
    competencies = [
        'Mobile Development: Flutter (Dart), Android (Java/Kotlin), iOS, Cross-Platform',
        'Backend Development: ASP.NET MVC, .NET Core, C#, REST APIs, GraphQL, Node.js',
        'AI/ML & Computer Vision: InsightFace, OCR (Qwen2-7B), TensorFlow Lite, Facial Recognition',
        'Biometric Authentication: BioPass Fingerprint SDK, Voice Recognition, Multi-Factor Auth',
        'AI-Powered Development: Cursor AI, GitHub Copilot, Vibe Coding, Prompt Engineering',
        'Architecture & Patterns: MVVM, Clean Architecture, Microservices, BLoC, Provider, GetX',
        'Real-Time Systems: XMPP (Ejabberd, xmpp_plugin), WebSocket, SignalR, Push Notifications',
        'Blockchain & Crypto: Node-based Blockchain, Cryptocurrency Wallet Integration',
        'Databases: SQL Server, PostgreSQL, MySQL, SQLite, Firebase, MongoDB',
        'Admin Systems: Admin Panels, User Management, Analytics Dashboards, RBAC',
        'State Management: BLoC, Provider, GetX, Riverpod',
        'DevOps & Tools: CI/CD Pipelines, Git, Docker, Automated Testing, Performance Optimization'
    ]
    for comp in competencies:
        p = doc.add_paragraph(comp, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    
    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    # Flutter Consultant
    exp1_header = doc.add_paragraph()
    exp1_header.add_run('Flutter Consultant').font.bold = True
    exp1_header.add_run('  |  ').font.bold = False
    exp1_header.add_run('Swati Technologies | Lahore, Pakistan').font.italic = True
    exp1_header.add_run('  |  ').font.italic = False
    exp1_header.add_run('October 2023 - Present').font.size = Pt(10)
    
    exp1_items = [
        'Leading development of cutting-edge mobile applications with AI, blockchain, and secure communication technologies, managing complete SDLC with 100% on-time delivery rate',
        'Pioneering AI-augmented development workflows using Cursor AI, GitHub Copilot, and vibe coding techniques, improving team productivity by 45%',
        'Architected and developed D-iDconnect Ecosystem: Revolutionary AI-powered security application serving thousands of users with complete administrative control system',
        'Integrated BioPass Fingerprint SDK (Flutter frontend kit) with .NET Core backend biometric verification models for secure multi-factor authentication',
        'Implemented facial recognition using InsightFace models (SCRFD-10GF detection, ResNet50 recognition) with liveness detection and anti-spoofing mechanisms',
        'Developed OCR document verification system using Qwen2-7B-Instruct model for automated identity document scanning, data extraction, and fraud detection',
        'Built complete admin panel system with comprehensive dashboard, user management, analytics, RBAC, audit logs, and real-time monitoring using SignalR',
        'Architected microservices backend with .NET Core and Node.js handling 100K+ daily transactions with optimized performance and scalability',
        'Implemented secure real-time XMPP chat system using Ejabberd server with xmpp_plugin 2.2.13, featuring end-to-end encryption, presence management, and offline message queuing',
        'Integrated node-based blockchain for secure transaction verification and cryptocurrency wallet management with mnemonic phrase generation',
        'Utilized AI-powered development tools reducing development time by 40% while maintaining 95% code quality through automated testing achieving 90% coverage',
        'Led 8+ developer team delivering enterprise mobile and backend solutions with modern AI-powered workflows'
    ]
    for item in exp1_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    
    featured = doc.add_paragraph()
    featured.add_run('🌟 Featured Project: D-iDconnect').font.bold = True
    featured.add_run(' — Complete ecosystem including mobile app (Flutter), admin panel (ASP.NET MVC), .NET Core APIs, blockchain integration, and AI-powered features. ')
    featured.add_run('Portfolio: portfolio-details-didconnect.html').font.italic = True
    featured.add_run(' | ')
    featured.add_run('Google Play').font.italic = True
    featured.add_run(' | ')
    featured.add_run('Apple App Store').font.italic = True
    featured.paragraph_format.space_after = Pt(12)
    
    # Team Lead
    exp2_header = doc.add_paragraph()
    exp2_header.add_run('Team Lead').font.bold = True
    exp2_header.add_run('  |  ').font.bold = False
    exp2_header.add_run('United Soft Labs | Lahore, Pakistan').font.italic = True
    exp2_header.add_run('  |  ').font.italic = False
    exp2_header.add_run('November 2020 - October 2023').font.size = Pt(10)
    
    exp2_items = [
        'Led 8+ developer team delivering enterprise mobile and backend solutions, established coding standards, CI/CD pipelines, and mentoring programs achieving 95% project success rate',
        'Developed RESTful APIs using ASP.NET MVC and .NET Core for seamless mobile-backend integration with complex SQL Server databases, stored procedures, and optimized queries',
        'Reduced bug reports by 60% through comprehensive code review standards and improved test coverage from 40% to 85% with automated testing frameworks',
        'Mentored 5 junior developers with 3 promoted to mid-level positions through technical guidance and professional development programs',
        'Managed complete SDLC from requirements gathering to deployment, ensuring on-time delivery and high-quality code standards',
        'Key Projects: aCart Online Shopping & Seller apps, Med Soft (Pharmacy POS), Dr MAN Pharmacy management system'
    ]
    for item in exp2_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph('Key Projects: portfolio-details-acart-onlineshopping.html | portfolio-details-acart-seller.html | portfolio-details-med-soft.html', style='List Bullet')
    p.paragraph_format.space_after = Pt(12)
    
    # Senior Android Developer
    exp3_header = doc.add_paragraph()
    exp3_header.add_run('Senior Android Developer').font.bold = True
    exp3_header.add_run('  |  ').font.bold = False
    exp3_header.add_run('Convergence Business Systems | Islamabad, Pakistan').font.italic = True
    exp3_header.add_run('  |  ').font.italic = False
    exp3_header.add_run('June 2019 - October 2020').font.size = Pt(10)
    
    exp3_items = [
        'Developed high-performance native Android and Flutter applications with .NET backend integration, transitioning legacy applications to Flutter reducing codebase by 40%',
        'Optimized performance reducing memory usage by 30% and improved user engagement by 35% through custom animations and enhanced UI/UX',
        'Integrated RESTful APIs for real-time data synchronization, conducted unit and widget testing ensuring robust application functionality',
        'Designed and implemented robust architectures for Android apps focusing on scalability and maintainability using Java and Kotlin',
        'Researched and implemented new Flutter libraries and technologies to improve development efficiency'
    ]
    for item in exp3_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()  # Spacing
    
    # Android & .NET Developer (Exceleron)
    exp4_header = doc.add_paragraph()
    exp4_header.add_run('Android & .NET Developer').font.bold = True
    exp4_header.add_run('  |  ').font.bold = False
    exp4_header.add_run('Exceleron Communications | Lahore, Pakistan').font.italic = True
    exp4_header.add_run('  |  ').font.italic = False
    exp4_header.add_run('December 2016 - June 2019').font.size = Pt(10)
    
    exp4_items = [
        'Delivered full-stack mobile and web solutions for enterprise clients including PTCL and Jazz Telecom',
        'Developed web applications using ASP.NET MVC, C#, and JavaScript with database schemas and RESTful APIs',
        'PTCL Smart Link Chat: Real-time messaging with XMPP integration serving 10K+ users',
        'Jazz Preventive Maintenance: Field worker app with GPS tracking and offline sync capabilities',
        'Designed and implemented database schemas, stored procedures, and optimized queries for high-performance applications',
        'Collaborated with cross-functional teams to gather and analyze requirements, translating them into technical specifications and solutions'
    ]
    for item in exp4_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()  # Spacing
    
    # Android & .NET Developer (Radius)
    exp5_header = doc.add_paragraph()
    exp5_header.add_run('Android & .NET Developer').font.bold = True
    exp5_header.add_run('  |  ').font.bold = False
    exp5_header.add_run('Radius Interactive | Pakistan').font.italic = True
    exp5_header.add_run('  |  ').font.italic = False
    exp5_header.add_run('November 2015 - November 2016').font.size = Pt(10)
    
    exp5_items = [
        'Developed web and mobile applications using ASP.NET, C#, and Android (Java)',
        'Integrated mobile applications with RESTful APIs for seamless data exchange',
        'Resolved complex software issues ensuring optimal performance and implemented version control workflows'
    ]
    for item in exp5_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()  # Spacing
    
    # Key Technical Implementations
    doc.add_heading('KEY TECHNICAL IMPLEMENTATIONS', level=1)
    tech_items = [
        'Face Recognition: InsightFace (SCRFD-10GF detection, ResNet50 recognition) with liveness detection and anti-spoofing - GitHub: github.com/deepinsight/insightface',
        'BioPass Fingerprint Biometric: BioPass ID Fingerprint SDK (Flutter) integrated with .NET Core backend models - biopassid.com/sdk/fingerprintsdk',
        'OCR Document Verification: Qwen2-7B-Instruct model for ID scanning and fraud detection - GitHub: github.com/yangjianxin1/Qwen2',
        'Admin Panel System: Full-featured admin dashboard with user management, analytics, RBAC, audit logs, and reporting',
        'Blockchain Integration: Node-based blockchain implementation for secure transactions and cryptocurrency wallet management',
        'Real-Time XMPP Chat: xmpp_plugin 2.2.13 with Ejabberd backend featuring end-to-end encryption and offline message queuing',
        '.NET Core Microservices: Scalable backend architecture with SQL Server RDBMS handling 100K+ daily transactions',
        'Real-Time Systems: WebSocket, SignalR, and XMPP for instant messaging and live admin updates',
        'AI-Assisted Testing: Automated test generation and quality assurance using AI tools achieving 90% code coverage'
    ]
    for item in tech_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()  # Spacing
    
    # Key Projects
    doc.add_heading('KEY PROJECTS', level=1)
    
    # D-iDconnect
    proj1 = doc.add_paragraph()
    proj1.add_run('🌟 D-iDconnect — AI-Powered Decentralized Identity Platform').font.bold = True
    p = doc.add_paragraph('Revolutionary AI-powered security application with multi-factor biometric authentication (face, fingerprint, voice), blockchain wallets, real-time XMPP chat, complete admin panel, and consent-based identity management. Available on Google Play and Apple App Store.')
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph('Portfolio: portfolio-details-didconnect.html | Google Play Store | Apple App Store')
    p2.paragraph_format.space_after = Pt(12)
    
    # aCart
    proj2 = doc.add_paragraph()
    proj2.add_run('aCart Online Shopping & Seller — E-commerce Marketplace').font.bold = True
    p = doc.add_paragraph('Complete marketplace ecosystem with customer app, seller app, and admin panel. Features include product management, order processing, payment integration, delivery tracking, and analytics dashboard.')
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph('Customer App: portfolio-details-acart-onlineshopping.html | Seller App: portfolio-details-acart-seller.html')
    p2.paragraph_format.space_after = Pt(12)
    
    # Eezly
    proj3 = doc.add_paragraph()
    proj3.add_run('Eezly — Food Delivery Application (Canada)').font.bold = True
    p = doc.add_paragraph('Cross-platform food delivery app with real-time order tracking, payment integration, and restaurant management system.')
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph('Portfolio: portfolio-details-eezly-food-app.html | Google Play | Apple App Store')
    p2.paragraph_format.space_after = Pt(12)
    
    # Med Soft
    proj5 = doc.add_paragraph()
    proj5.add_run('Med Soft — Healthcare Management System').font.bold = True
    p = doc.add_paragraph('Comprehensive pharmacy POS system with inventory management, prescription processing, patient records, sales, purchases, and reporting. Features ChatGPT integration and multi-platform support.')
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph('Portfolio: portfolio-details-med-soft.html')
    p2.paragraph_format.space_after = Pt(12)
    
    # AI-Powered Development
    doc.add_heading('AI-POWERED DEVELOPMENT EXPERTISE', level=1)
    ai_items = [
        'Cursor AI Integration: Leveraging Cursor AI for intelligent code completion, refactoring, and bug detection across Flutter and .NET projects, reducing development time by 40%',
        'GitHub Copilot: Utilizing AI-powered code reviews and suggestions maintaining 95% code quality standards with automated test generation achieving 90% coverage',
        'Vibe Coding & Prompt Engineering: Mastered vibe coding techniques for rapid prototyping, expert in crafting precise prompts for code generation, architecture design, and problem-solving',
        'AI-Assisted Development Practices: Implementing AI-assisted test generation, documentation generation, code explanations, technical debt analysis, and performance optimization suggestions',
        'Team Training: Created comprehensive prompt engineering guidelines and training programs for team members on effective AI tool utilization'
    ]
    for item in ai_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()  # Spacing
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    edu_header = doc.add_paragraph()
    edu_header.add_run('Bachelor of Computer Science').font.bold = True
    edu_header.add_run('  |  ').font.bold = False
    edu_header.add_run('2010 - 2014').font.size = Pt(10)
    p = doc.add_paragraph('Virtual University of Pakistan', style='Intense Quote')
    p2 = doc.add_paragraph('Focused on software development fundamentals, algorithms, data structures, database systems, and distributed computing.')
    p2.paragraph_format.space_after = Pt(12)
    
    # Additional Information
    doc.add_heading('ADDITIONAL INFORMATION', level=1)
    add_items = [
        'Portfolio Website: info-abubakar-hussain.github.io/abubakarhussain',
        'LinkedIn: linkedin.com/in/abubakar-hussain-66b834130',
        'Freelance: Available for consulting and project work',
        'Languages: English (Fluent), Urdu (Native)',
        '10+ years professional software development experience',
        '5+ years Flutter cross-platform development',
        '4+ years Android native development'
    ]
    for item in add_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
    
    # Save document
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_file = os.path.join(script_dir, 'AbuBakar_Hussain_Resume.docx')
    doc.save(output_file)
    print(f"✅ Success! Word document created: {output_file}")
    print(f"📄 You can now open and edit this file in Microsoft Word")
    
except ImportError:
    print("❌ python-docx not installed.")
    print("\nTo install python-docx, run:")
    print("  pip3 install python-docx")
    print("\nThen run this script again:")
    print("  python3 create_resume_docx.py")
    
except Exception as e:
    print(f"❌ Error: {e}")
    import traceback
    traceback.print_exc()
